from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.schemas.study_report import StudyReportAnalysis


# =========================================================
# 경로
# =========================================================

BASE_DIR = Path(__file__).resolve().parent

REPORT_DIR = (
    BASE_DIR
    / "ai-interview-report"
)

INPUT_PATH = (
    REPORT_DIR
    / "study_report.json"
)

OUTPUT_PATH = (
    REPORT_DIR
    / "study_report.docx"
)


# =========================================================
# 색상
# =========================================================

COLOR_PRIMARY = "1F4E78"
COLOR_SECONDARY = "D9EAF7"
COLOR_LIGHT = "F5F7FA"
COLOR_BORDER = "D9DEE5"

COLOR_GREEN_LIGHT = "E2F0D9"
COLOR_YELLOW_LIGHT = "FFF2CC"
COLOR_RED_LIGHT = "FDE9E7"

COLOR_HIGH = "C0504D"
COLOR_MEDIUM = "BF9000"
COLOR_LOW = "548235"

COLOR_TEXT = RGBColor(
    40,
    40,
    40,
)

COLOR_MUTED = RGBColor(
    105,
    105,
    105,
)

COLOR_PRIMARY_RGB = RGBColor(
    31,
    78,
    120,
)


# =========================================================
# 표시 이름
# =========================================================

DISPLAY_MAP = {
    "high": "높음",
    "medium": "중간",
    "low": "낮음",
    "not_covered": "미확보",

    "strong": "강함",
    "moderate": "중간",
    "weak": "약함",

    "explicit_user_request": (
        "참여자 직접 요구"
    ),

    "derived_opportunity": (
        "분석 기반 기회"
    ),
}


# =========================================================
# JSON 로드
# =========================================================

def load_study_report(
) -> StudyReportAnalysis:

    if not INPUT_PATH.exists():

        raise FileNotFoundError(
            "study_report.json 파일을 "
            "찾을 수 없습니다.\n"
            f"경로: {INPUT_PATH}"
        )

    with INPUT_PATH.open(
        "r",
        encoding="utf-8",
    ) as file:

        raw_data = json.load(
            file
        )

    return (
        StudyReportAnalysis
        .model_validate(
            raw_data
        )
    )


# =========================================================
# 공통 Helper
# =========================================================

def display_value(
    value: str | None,
) -> str:

    if value is None:
        return "-"

    return DISPLAY_MAP.get(
        value,
        value,
    )


def hex_to_rgb(
    value: str,
) -> RGBColor:

    value = value.lstrip("#")

    return RGBColor(
        int(value[0:2], 16),
        int(value[2:4], 16),
        int(value[4:6], 16),
    )


def importance_color(
    value: str | None,
) -> RGBColor:

    if value in (
        "high",
        "strong",
    ):
        return hex_to_rgb(
            COLOR_HIGH
        )

    if value in (
        "medium",
        "moderate",
    ):
        return hex_to_rgb(
            COLOR_MEDIUM
        )

    return hex_to_rgb(
        COLOR_LOW
    )


def importance_fill(
    value: str | None,
) -> str:

    if value in (
        "high",
        "strong",
    ):
        return COLOR_RED_LIGHT

    if value in (
        "medium",
        "moderate",
    ):
        return COLOR_YELLOW_LIGHT

    return COLOR_GREEN_LIGHT


def set_cell_shading(
    cell,
    fill: str,
) -> None:

    tc_pr = (
        cell._tc
        .get_or_add_tcPr()
    )

    shading = OxmlElement(
        "w:shd"
    )

    shading.set(
        qn("w:fill"),
        fill,
    )

    tc_pr.append(
        shading
    )


def set_cell_margins(
    cell,
    top: int = 80,
    start: int = 100,
    bottom: int = 80,
    end: int = 100,
) -> None:

    tc = cell._tc

    tc_pr = (
        tc.get_or_add_tcPr()
    )

    tc_mar = (
        tc_pr.first_child_found_in(
            "w:tcMar"
        )
    )

    if tc_mar is None:

        tc_mar = OxmlElement(
            "w:tcMar"
        )

        tc_pr.append(
            tc_mar
        )

    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):

        node = tc_mar.find(
            qn(
                f"w:{margin_name}"
            )
        )

        if node is None:

            node = OxmlElement(
                f"w:{margin_name}"
            )

            tc_mar.append(
                node
            )

        node.set(
            qn("w:w"),
            str(value),
        )

        node.set(
            qn("w:type"),
            "dxa",
        )


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    font_size: float = 9,
    color: RGBColor | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:

    cell.text = ""

    paragraph = (
        cell.paragraphs[0]
    )

    if align is not None:
        paragraph.alignment = align

    run = paragraph.add_run(
        str(text)
    )

    run.bold = bold

    run.font.name = (
        "Malgun Gothic"
    )

    run.font.size = Pt(
        font_size
    )

    if color is not None:
        run.font.color.rgb = color

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT
        .CENTER
    )

    set_cell_margins(
        cell
    )


def set_table_borders(
    table,
    color: str = COLOR_BORDER,
    size: int = 4,
) -> None:

    tbl_pr = (
        table._tbl.tblPr
    )

    borders = (
        tbl_pr.first_child_found_in(
            "w:tblBorders"
        )
    )

    if borders is None:

        borders = OxmlElement(
            "w:tblBorders"
        )

        tbl_pr.append(
            borders
        )

    for edge in (
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ):

        tag = f"w:{edge}"

        element = borders.find(
            qn(tag)
        )

        if element is None:

            element = OxmlElement(
                tag
            )

            borders.append(
                element
            )

        element.set(
            qn("w:val"),
            "single",
        )

        element.set(
            qn("w:sz"),
            str(size),
        )

        element.set(
            qn("w:space"),
            "0",
        )

        element.set(
            qn("w:color"),
            color,
        )


def add_table_header(
    table,
    headers: list[str],
) -> None:

    row = table.rows[0]

    for index, header in enumerate(
        headers
    ):

        cell = row.cells[index]

        set_cell_shading(
            cell,
            COLOR_PRIMARY,
        )

        set_cell_text(
            cell,
            header,
            bold=True,
            font_size=8.5,
            color=RGBColor(
                255,
                255,
                255,
            ),
        )


# =========================================================
# 문서 텍스트
# =========================================================

def add_section_heading(
    document: Document,
    number: str,
    title: str,
) -> None:

    paragraph = (
        document.add_paragraph()
    )

    paragraph.paragraph_format.space_before = Pt(
        16
    )

    paragraph.paragraph_format.space_after = Pt(
        7
    )

    paragraph.paragraph_format.keep_with_next = True

    number_run = (
        paragraph.add_run(
            f"{number} "
        )
    )

    number_run.bold = True

    number_run.font.name = (
        "Malgun Gothic"
    )

    number_run.font.size = Pt(
        17
    )

    number_run.font.color.rgb = (
        COLOR_PRIMARY_RGB
    )

    title_run = (
        paragraph.add_run(
            title
        )
    )

    title_run.bold = True

    title_run.font.name = (
        "Malgun Gothic"
    )

    title_run.font.size = Pt(
        17
    )

    title_run.font.color.rgb = (
        COLOR_TEXT
    )


def add_subheading(
    document: Document,
    title: str,
) -> None:

    paragraph = (
        document.add_paragraph()
    )

    paragraph.paragraph_format.space_before = Pt(
        8
    )

    paragraph.paragraph_format.space_after = Pt(
        3
    )

    paragraph.paragraph_format.keep_with_next = True

    run = paragraph.add_run(
        title
    )

    run.bold = True

    run.font.name = (
        "Malgun Gothic"
    )

    run.font.size = Pt(
        11
    )

    run.font.color.rgb = (
        COLOR_PRIMARY_RGB
    )


def add_body(
    document: Document,
    text: str,
    *,
    bold_prefix: str | None = None,
) -> None:

    paragraph = (
        document.add_paragraph()
    )

    paragraph.paragraph_format.space_after = Pt(
        5
    )

    paragraph.paragraph_format.line_spacing = 1.2

    if bold_prefix:

        prefix_run = (
            paragraph.add_run(
                bold_prefix
            )
        )

        prefix_run.bold = True

        prefix_run.font.name = (
            "Malgun Gothic"
        )

        prefix_run.font.size = Pt(
            9.5
        )

    run = paragraph.add_run(
        text
    )

    run.font.name = (
        "Malgun Gothic"
    )

    run.font.size = Pt(
        9.5
    )

    run.font.color.rgb = (
        COLOR_TEXT
    )


def add_meta_line(
    document: Document,
    *,
    label: str | None = None,
    status: str | None = None,
    participant_count: int | None = None,
    extra_text: str | None = None,
) -> None:

    paragraph = (
        document.add_paragraph()
    )

    paragraph.paragraph_format.space_after = Pt(
        3
    )

    if status:

        dot = paragraph.add_run(
            "● "
        )

        dot.font.name = (
            "Malgun Gothic"
        )

        dot.font.size = Pt(
            9
        )

        dot.font.color.rgb = (
            importance_color(
                status
            )
        )

    if label and status:

        run = paragraph.add_run(
            (
                f"{label} "
                f"{display_value(status)}"
            )
        )

        run.bold = True

        run.font.name = (
            "Malgun Gothic"
        )

        run.font.size = Pt(
            8.5
        )

        run.font.color.rgb = (
            COLOR_TEXT
        )

    if participant_count is not None:

        if label and status:

            separator = paragraph.add_run(
                "   |   "
            )

            separator.font.color.rgb = (
                COLOR_MUTED
            )

        run = paragraph.add_run(
            (
                f"참여자 "
                f"{participant_count}명"
            )
        )

        run.font.name = (
            "Malgun Gothic"
        )

        run.font.size = Pt(
            8.5
        )

        run.font.color.rgb = (
            COLOR_MUTED
        )

    if extra_text:

        separator = paragraph.add_run(
            "   |   "
        )

        separator.font.color.rgb = (
            COLOR_MUTED
        )

        run = paragraph.add_run(
            extra_text
        )

        run.font.name = (
            "Malgun Gothic"
        )

        run.font.size = Pt(
            8.5
        )

        run.font.color.rgb = (
            COLOR_MUTED
        )


def add_info_box(
    document: Document,
    title: str,
    text: str,
) -> None:

    table = document.add_table(
        rows=1,
        cols=1,
    )

    cell = table.cell(
        0,
        0,
    )

    set_cell_shading(
        cell,
        COLOR_SECONDARY,
    )

    paragraph = (
        cell.paragraphs[0]
    )

    title_run = (
        paragraph.add_run(
            f"{title}\n"
        )
    )

    title_run.bold = True

    title_run.font.name = (
        "Malgun Gothic"
    )

    title_run.font.size = Pt(
        10
    )

    title_run.font.color.rgb = (
        COLOR_PRIMARY_RGB
    )

    body_run = (
        paragraph.add_run(
            text
        )
    )

    body_run.font.name = (
        "Malgun Gothic"
    )

    body_run.font.size = Pt(
        9.5
    )

    body_run.font.color.rgb = (
        COLOR_TEXT
    )

    set_cell_margins(
        cell,
        top=160,
        start=180,
        bottom=160,
        end=180,
    )

    set_table_borders(
        table,
        color=COLOR_SECONDARY,
    )

    document.add_paragraph()


# =========================================================
# Evidence
# =========================================================

def build_evidence_map(
    report: StudyReportAnalysis,
) -> dict[str, Any]:

    return {
        item.evidence_id: item
        for item in report.evidence
    }


def add_representative_evidence(
    document: Document,
    evidence_ids: list[str],
    evidence_map: dict[str, Any],
) -> None:

    representative = None

    for evidence_id in evidence_ids:

        candidate = (
            evidence_map.get(
                evidence_id
            )
        )

        if candidate is not None:

            representative = (
                candidate
            )

            break

    if representative is None:
        return

    table = document.add_table(
        rows=1,
        cols=1,
    )

    cell = table.cell(
        0,
        0,
    )

    set_cell_shading(
        cell,
        COLOR_LIGHT,
    )

    paragraph = (
        cell.paragraphs[0]
    )

    label_run = (
        paragraph.add_run(
            "대표 발언\n"
        )
    )

    label_run.bold = True

    label_run.font.name = (
        "Malgun Gothic"
    )

    label_run.font.size = Pt(
        8
    )

    label_run.font.color.rgb = (
        COLOR_PRIMARY_RGB
    )

    quote_run = (
        paragraph.add_run(
            (
                f"“{representative.quote}”"
                f"\n- "
                f"{representative.participant_id}"
            )
        )
    )

    quote_run.font.name = (
        "Malgun Gothic"
    )

    quote_run.font.size = Pt(
        8.5
    )

    quote_run.italic = True

    quote_run.font.color.rgb = (
        COLOR_TEXT
    )

    set_cell_margins(
        cell,
        top=120,
        start=160,
        bottom=120,
        end=160,
    )

    set_table_borders(
        table
    )

    document.add_paragraph()


# =========================================================
# 문서 설정
# =========================================================

def configure_document(
    document: Document,
) -> None:

    section = (
        document.sections[0]
    )

    section.top_margin = Cm(
        1.8
    )

    section.bottom_margin = Cm(
        1.8
    )

    section.left_margin = Cm(
        1.8
    )

    section.right_margin = Cm(
        1.8
    )

    normal = (
        document.styles[
            "Normal"
        ]
    )

    normal.font.name = (
        "Malgun Gothic"
    )

    normal.font.size = Pt(
        9.5
    )

    normal._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Malgun Gothic",
    )


# =========================================================
# 표지
# =========================================================

def add_cover(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    overview = report.overview

    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run("AI INTERVIEW")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_PRIMARY_RGB

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(12)

    run = paragraph.add_run("Research Report")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(30)
    run.font.color.rgb = COLOR_TEXT

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(overview.research_title)
    run.bold = True
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(17)
    run.font.color.rgb = COLOR_PRIMARY_RGB

    document.add_paragraph()

    # 질문 수는 실제 인터뷰의 총 질문 횟수와 혼동될 수 있으므로
    # 표지에서는 제외한다.
    table = document.add_table(
        rows=3,
        cols=2,
    )

    labels = [
        ("Study ID", overview.study_id),
        ("참여자", f"{overview.participant_count}명"),
        ("완료 인터뷰", f"{overview.completed_session_count}건"),
    ]

    for row_index, (label, value) in enumerate(labels):
        left = table.cell(row_index, 0)
        right = table.cell(row_index, 1)

        set_cell_shading(left, COLOR_SECONDARY)

        set_cell_text(
            left,
            label,
            bold=True,
            color=COLOR_PRIMARY_RGB,
        )

        set_cell_text(
            right,
            value,
        )

    set_table_borders(table)

    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        "Research Purpose\n"
        f"{overview.research_purpose}"
    )
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEXT

    document.add_paragraph()
    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        "Generated "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    document.add_page_break()


# =========================================================
# Research Snapshot
# =========================================================

def build_bar(
    value: int,
    maximum: int,
    *,
    width: int = 10,
) -> str:

    if maximum <= 0:
        return "·" * width

    filled = round(
        (
            value
            / maximum
        )
        * width
    )

    if (
        value > 0
        and filled == 0
    ):
        filled = 1

    filled = max(
        0,
        min(
            filled,
            width,
        ),
    )

    return (
        "■" * filled
        + "·" * (
            width - filled
        )
    )


def add_metric_card(
    cell,
    label: str,
    value: str,
    *,
    accent: RGBColor | None = None,
) -> None:

    set_cell_shading(
        cell,
        COLOR_LIGHT,
    )

    cell.text = ""

    paragraph = (
        cell.paragraphs[0]
    )

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    value_run = (
        paragraph.add_run(
            f"{value}\n"
        )
    )

    value_run.bold = True

    value_run.font.name = (
        "Malgun Gothic"
    )

    value_run.font.size = Pt(
        19
    )

    value_run.font.color.rgb = (
        accent
        if accent is not None
        else COLOR_PRIMARY_RGB
    )

    label_run = (
        paragraph.add_run(
            label
        )
    )

    label_run.font.name = (
        "Malgun Gothic"
    )

    label_run.font.size = Pt(
        8
    )

    label_run.font.color.rgb = (
        COLOR_MUTED
    )

    set_cell_margins(
        cell,
        top=180,
        start=80,
        bottom=180,
        end=80,
    )


def add_snapshot_rank_table(
    document: Document,
    title: str,
    items: list[Any],
    *,
    participant_total: int,
    title_getter,
    count_getter,
) -> None:

    add_subheading(
        document,
        title,
    )

    if not items:
        add_body(
            document,
            "해당 분석 결과가 없습니다.",
        )
        return

    table = document.add_table(
        rows=1,
        cols=2,
    )

    add_table_header(
        table,
        [
            "항목",
            "Evidence Participants",
        ],
    )

    for item in items:
        count = count_getter(item)

        row = table.add_row()

        set_cell_text(
            row.cells[0],
            title_getter(item),
            font_size=8,
        )

        cell = row.cells[1]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        bar_run = paragraph.add_run(
            build_bar(
                count,
                participant_total,
                width=12,
            )
        )
        bar_run.bold = True
        bar_run.font.name = "Malgun Gothic"
        bar_run.font.size = Pt(10)
        bar_run.font.color.rgb = COLOR_PRIMARY_RGB

        count_run = paragraph.add_run(
            f"\n{count} / {participant_total}명"
        )
        count_run.bold = True
        count_run.font.name = "Malgun Gothic"
        count_run.font.size = Pt(8.5)
        count_run.font.color.rgb = COLOR_TEXT

    set_table_borders(table)



def add_research_snapshot(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    overview = report.overview

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)

    run = paragraph.add_run("RESEARCH SNAPSHOT")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(21)
    run.font.color.rgb = COLOR_PRIMARY_RGB

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)

    run = paragraph.add_run(
        "의사결정에 필요한 핵심 결과만 요약했습니다."
    )
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED

    # -----------------------------------------------------
    # Study At A Glance
    # 실제 해석 가치가 있는 지표만 표시한다.
    # -----------------------------------------------------

    high_priority_opportunities = sum(
        1
        for item in report.opportunities
        if item.priority == "high"
    )

    metric_table = document.add_table(
        rows=1,
        cols=4,
    )

    add_metric_card(
        metric_table.cell(0, 0),
        "참여자",
        str(overview.participant_count),
    )

    add_metric_card(
        metric_table.cell(0, 1),
        "완료 인터뷰",
        str(overview.completed_session_count),
    )

    add_metric_card(
        metric_table.cell(0, 2),
        "전체 Coverage",
        display_value(
            report.research_coverage.overall_coverage
        ),
        accent=importance_color(
            report.research_coverage.overall_coverage
        ),
    )

    add_metric_card(
        metric_table.cell(0, 3),
        "High Priority Opportunities",
        str(high_priority_opportunities),
        accent=importance_color("high"),
    )

    set_table_borders(metric_table)

    document.add_paragraph()

    # -----------------------------------------------------
    # Top Findings
    # 막대 기준은 항상 전체 참여자 수다.
    # -----------------------------------------------------

    top_findings = sorted(
        report.key_findings,
        key=lambda item: item.participant_count,
        reverse=True,
    )[:3]

    add_snapshot_rank_table(
        document,
        "Top Key Findings",
        top_findings,
        participant_total=overview.participant_count,
        title_getter=lambda item: item.title,
        count_getter=lambda item: item.participant_count,
    )

    # -----------------------------------------------------
    # Top Pain Points
    # -----------------------------------------------------

    top_pains = sorted(
        report.pain_points,
        key=lambda item: item.participant_count,
        reverse=True,
    )[:3]

    add_snapshot_rank_table(
        document,
        "Top Pain Points",
        top_pains,
        participant_total=overview.participant_count,
        title_getter=lambda item: item.title,
        count_getter=lambda item: item.participant_count,
    )

    # -----------------------------------------------------
    # 해석 주의 문구
    # 정성 인터뷰의 참여자 수를 모집단 비율처럼 해석하지 않도록 명시한다.
    # -----------------------------------------------------

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)

    run = paragraph.add_run(
        "※ Evidence Participants는 해당 인사이트를 뒷받침하는 "
        "직접 Evidence가 확인된 참여자 수입니다. "
        "모집단 비율이나 시장 점유율을 의미하지 않습니다."
    )
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(7.5)
    run.font.color.rgb = COLOR_MUTED

    # Snapshot은 여기서 끝낸다.
    # Opportunity Priority 표는 High Priority Opportunities 카드와
    # 내용이 중복되고 페이지를 불필요하게 차지하므로 제외한다.
    document.add_page_break()


# =========================================================
# 01 Executive Summary
# =========================================================

def add_executive_summary(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    summary = (
        report.executive_summary
    )

    add_section_heading(
        document,
        "01",
        "Executive Summary",
    )

    add_info_box(
        document,
        "Core Insight",
        summary.core_insight,
    )

    add_body(
        document,
        summary.summary,
    )

    add_subheading(
        document,
        "Key Takeaways",
    )

    for index, item in enumerate(
        summary.key_takeaways,
        start=1,
    ):

        paragraph = (
            document.add_paragraph()
        )

        paragraph.paragraph_format.space_after = Pt(
            2
        )

        run = paragraph.add_run(
            (
                f"{index}. "
                f"{item.point}"
            )
        )

        run.bold = True

        run.font.name = (
            "Malgun Gothic"
        )

        run.font.size = Pt(
            9.5
        )

        run.font.color.rgb = (
            COLOR_TEXT
        )

        add_meta_line(
            document,
            participant_count=(
                item.participant_count
            ),
        )


# =========================================================
# 02 Research Coverage
# =========================================================

def add_research_coverage(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    coverage = (
        report.research_coverage
    )

    add_section_heading(
        document,
        "02",
        "Research Coverage",
    )

    add_info_box(
        document,
        "Overall Coverage",
        display_value(
            coverage.overall_coverage
        ),
    )

    table = document.add_table(
        rows=1,
        cols=4,
    )

    add_table_header(
        table,
        [
            "Question",
            "Coverage",
            "응답",
            "분석",
        ],
    )

    for item in coverage.items:

        row = (
            table.add_row()
        )

        cells = (
            row.cells
        )

        set_cell_text(
            cells[0],
            (
                f"{item.question_id}\n"
                f"{item.question}"
            ),
            font_size=8,
        )

        set_cell_text(
            cells[1],
            display_value(
                item.coverage
            ),
            bold=True,
            font_size=8,
        )

        set_cell_shading(
            cells[1],
            importance_fill(
                item.coverage
            ),
        )

        set_cell_text(
            cells[2],
            (
                f"{item.covered_participant_count}"
                f" / "
                f"{item.participant_count}"
            ),
            font_size=8,
            align=(
                WD_ALIGN_PARAGRAPH.CENTER
            ),
        )

        analysis_text = (
            item.reason
        )

        if item.missing_information:

            analysis_text += (
                "\n\n추가 확인 필요: "
                + " / ".join(
                    item.missing_information
                )
            )

        set_cell_text(
            cells[3],
            analysis_text,
            font_size=8,
        )

    set_table_borders(
        table
    )


# =========================================================
# 03 Key Findings
# =========================================================

def add_key_findings(
    document: Document,
    report: StudyReportAnalysis,
    evidence_map: dict[str, Any],
) -> None:

    add_section_heading(
        document,
        "03",
        "Key Findings",
    )

    for index, item in enumerate(
        report.key_findings,
        start=1,
    ):

        add_subheading(
            document,
            (
                f"{index}. "
                f"{item.title}"
            ),
        )

        add_meta_line(
            document,
            label="Evidence",
            status=(
                item.evidence_strength
            ),
            participant_count=(
                item.participant_count
            ),
        )

        add_body(
            document,
            item.summary,
        )

        add_representative_evidence(
            document,
            item.evidence_ids,
            evidence_map,
        )


# =========================================================
# 04 Themes
# =========================================================

def add_themes(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    add_section_heading(
        document,
        "04",
        "Themes",
    )

    for index, item in enumerate(
        report.themes,
        start=1,
    ):

        add_subheading(
            document,
            (
                f"{index}. "
                f"{item.theme}"
            ),
        )

        add_meta_line(
            document,
            participant_count=(
                item.participant_count
            ),
        )

        add_body(
            document,
            item.description,
        )


# =========================================================
# 05 Key Drivers
# =========================================================

def add_key_drivers(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    add_section_heading(
        document,
        "05",
        "Key Drivers",
    )

    for index, item in enumerate(
        report.key_drivers,
        start=1,
    ):

        add_subheading(
            document,
            (
                f"{index}. "
                f"{item.driver}"
            ),
        )

        add_meta_line(
            document,
            label="Strength",
            status=(
                item.strength
            ),
            participant_count=(
                item.participant_count
            ),
        )

        add_body(
            document,
            item.description,
        )


# =========================================================
# 06 Pain Points
# =========================================================

def add_pain_points(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    add_section_heading(
        document,
        "06",
        "Pain Points",
    )

    for index, item in enumerate(
        report.pain_points,
        start=1,
    ):

        add_subheading(
            document,
            (
                f"{index}. "
                f"{item.title}"
            ),
        )

        add_meta_line(
            document,
            label="Severity",
            status=(
                item.severity
            ),
            participant_count=(
                item.participant_count
            ),
        )

        add_body(
            document,
            item.description,
        )


# =========================================================
# 07 Needs
# =========================================================

def add_needs(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    add_section_heading(
        document,
        "07",
        "Needs",
    )

    for index, item in enumerate(
        report.needs,
        start=1,
    ):

        add_subheading(
            document,
            (
                f"{index}. "
                f"{item.title}"
            ),
        )

        add_meta_line(
            document,
            label="Priority",
            status=(
                item.priority
            ),
            participant_count=(
                item.participant_count
            ),
        )

        add_body(
            document,
            item.description,
        )


# =========================================================
# 08 Segment Differences
# =========================================================

def add_segments(
    document: Document,
    report: StudyReportAnalysis,
    evidence_map: dict[str, Any],
) -> None:

    add_section_heading(
        document,
        "08",
        "Segment Differences",
    )

    if not report.segment_differences:

        add_body(
            document,
            (
                "현재 Study에서는 충분한 근거를 가진 "
                "Segment Difference가 확인되지 않았습니다."
            ),
        )

        return

    for index, segment in enumerate(
        report.segment_differences,
        start=1,
    ):

        add_subheading(
            document,
            (
                f"{index}. "
                f"{segment.segment_name}"
            ),
        )

        add_meta_line(
            document,
            participant_count=(
                segment.participant_count
            ),
        )

        add_body(
            document,
            segment.segment_description,
        )

        add_info_box(
            document,
            "Key Difference",
            segment.key_difference,
        )

        for group_index, group in enumerate(
            segment.groups,
            start=1,
        ):

            add_subheading(
                document,
                (
                    f"Group {group_index}. "
                    f"{group.group_name}"
                ),
            )

            add_meta_line(
                document,
                participant_count=(
                    group.participant_count
                ),
            )

            add_body(
                document,
                group.group_description,
            )

            add_representative_evidence(
                document,
                group.evidence_ids,
                evidence_map,
            )


# =========================================================
# 09 Opportunities
# =========================================================

def add_opportunities(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    add_section_heading(
        document,
        "09",
        "Opportunities",
    )

    for index, item in enumerate(
        report.opportunities,
        start=1,
    ):

        add_subheading(
            document,
            (
                f"{index}. "
                f"{item.opportunity}"
            ),
        )

        add_meta_line(
            document,
            label="Priority",
            status=(
                item.priority
            ),
            participant_count=(
                item.participant_count
            ),
            extra_text=(
                display_value(
                    item.source_type
                )
            ),
        )

        add_body(
            document,
            item.problem_or_need,
            bold_prefix=(
                "Problem / Need: "
            ),
        )

        add_body(
            document,
            item.expected_value,
            bold_prefix=(
                "Expected Value: "
            ),
        )


# =========================================================
# 10 Research Gaps
# =========================================================

def add_research_gaps(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    add_section_heading(
        document,
        "10",
        "Research Gaps",
    )

    for index, item in enumerate(
        report.research_gaps,
        start=1,
    ):

        add_subheading(
            document,
            (
                f"{index}. "
                f"{item.topic}"
            ),
        )

        add_meta_line(
            document,
            label="Priority",
            status=(
                item.priority
            ),
        )

        add_body(
            document,
            item.reason,
        )


# =========================================================
# Footer
# =========================================================

def add_footer(
    document: Document,
    report: StudyReportAnalysis,
) -> None:

    for section in (
        document.sections
    ):

        footer = (
            section.footer
        )

        paragraph = (
            footer.paragraphs[0]
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            (
                "AI Interview Research Report"
                "  |  "
                f"{report.overview.study_id}"
            )
        )

        run.font.name = (
            "Malgun Gothic"
        )

        run.font.size = Pt(
            7
        )

        run.font.color.rgb = (
            COLOR_MUTED
        )


# =========================================================
# 문서 생성
# =========================================================

def build_document(
    report: StudyReportAnalysis,
) -> Document:

    document = Document()

    configure_document(
        document
    )

    evidence_map = (
        build_evidence_map(
            report
        )
    )

    # 표지
    add_cover(
        document,
        report,
    )

    # 새로 추가된 Snapshot
    add_research_snapshot(
        document,
        report,
    )

    add_executive_summary(
        document,
        report,
    )

    add_research_coverage(
        document,
        report,
    )

    add_key_findings(
        document,
        report,
        evidence_map,
    )

    add_themes(
        document,
        report,
    )

    add_key_drivers(
        document,
        report,
    )

    add_pain_points(
        document,
        report,
    )

    add_needs(
        document,
        report,
    )

    add_segments(
        document,
        report,
        evidence_map,
    )

    add_opportunities(
        document,
        report,
    )

    add_research_gaps(
        document,
        report,
    )

    add_footer(
        document,
        report,
    )

    return document


# =========================================================
# Main
# =========================================================

def main() -> None:

    print()

    print(
        "========================================"
    )

    print(
        "Study Report → Visual Summary Word"
    )

    print(
        "========================================"
    )

    report = (
        load_study_report()
    )

    document = (
        build_document(
            report
        )
    )

    REPORT_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        OUTPUT_PATH
    )

    print()

    print(
        "시각화 포함 Word 리포트 생성 완료"
    )

    print(
        f"Study: "
        f"{report.overview.study_id}"
    )

    print(
        f"참여자: "
        f"{report.overview.participant_count}명"
    )

    print(
        f"Key Findings: "
        f"{len(report.key_findings)}"
    )

    print(
        f"Pain Points: "
        f"{len(report.pain_points)}"
    )

    print(
        f"Opportunities: "
        f"{len(report.opportunities)}"
    )

    print()

    print(
        "Research Snapshot: 포함"
    )

    print(
        "Evidence Appendix: 제외"
    )

    print(
        "대표 Evidence: "
        "Key Findings / Segment만 표시"
    )

    print()

    print(
        f"생성 파일: "
        f"{OUTPUT_PATH}"
    )


if __name__ == "__main__":
    main()