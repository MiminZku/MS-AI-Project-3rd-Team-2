"""다운로드용 문서 생성 (인터뷰 기록 / 프로젝트 리포트).

홈페이지 다운로드 센터에서 클라이언트·PM이 바로 받아볼 수 있는 파일을 만든다.
python-docx가 없거나 실패하면 같은 내용을 텍스트로 떨어뜨려, 다운로드 자체가
막히는 일은 없게 한다.
"""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime
from typing import Any

from app.schemas.project_report import ProjectAggregateReport
from app.schemas.session import Session, Turn
from app.schemas.study import ResearchStudy

logger = logging.getLogger(__name__)

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_MEDIA_TYPE = "text/plain; charset=utf-8"

# 파일명에 쓸 수 없는 문자 (Windows 기준으로 넉넉히 제거)
_UNSAFE_FILENAME = re.compile(r'[\\/:*?"<>|\x00-\x1f]+')


def safe_filename_part(value: str, *, fallback: str = "untitled") -> str:
    cleaned = _UNSAFE_FILENAME.sub("_", (value or "").strip())
    cleaned = re.sub(r"\s+", "_", cleaned).strip("._")
    return cleaned or fallback


def _format_datetime(value: datetime | None) -> str:
    return value.strftime("%Y-%m-%d %H:%M") if value else "-"


# =========================================================
# 인터뷰 기록 (질문/답변)
# =========================================================

def _transcript_lines(
    session: Session,
    turns: list[Turn],
    *,
    project_title: str | None,
) -> list[tuple[str, str]]:
    """(스타일, 텍스트) 목록. docx와 txt 양쪽에서 같은 내용을 쓴다."""

    lines: list[tuple[str, str]] = [
        ("title", f"인터뷰 기록 — {session.title}"),
    ]

    if project_title:
        lines.append(("meta", f"프로젝트: {project_title}"))

    lines.extend(
        [
            ("meta", f"참가자 ID: {session.title}"),
            ("meta", f"세션 ID: {session.id}"),
            ("meta", f"상태: {session.status}"),
            ("meta", f"예정 시간: {session.duration_minutes}분"),
            ("meta", f"시작: {_format_datetime(session.started_at)}"),
            ("meta", f"종료: {_format_datetime(session.ended_at)}"),
            ("spacer", ""),
        ]
    )

    if not turns:
        lines.append(("body", "기록된 대화가 없습니다."))
        return lines

    for turn in turns:
        speaker = "AI 진행자" if turn.speaker == "assistant" else "응답자"
        lines.append(("speaker", f"[{turn.index}] {speaker}"))
        lines.append(("body", turn.text))
        if turn.text_en:
            lines.append(("translation", f"(EN) {turn.text_en}"))
        lines.append(("spacer", ""))

    return lines


def _lines_to_text(lines: list[tuple[str, str]]) -> bytes:
    out: list[str] = []
    for style, text in lines:
        if style == "spacer":
            out.append("")
        elif style == "title":
            out.append(text)
            out.append("=" * len(text))
        elif style == "heading":
            out.append("")
            out.append(f"## {text}")
        elif style == "speaker":
            out.append(f"{text}")
        else:
            out.append(text)
    return "\n".join(out).encode("utf-8")


# Word는 라틴 글꼴(기본 Calibri)만 지정된 문서에서 한글을 렌더링할 때
# East Asian 글꼴을 따로 보지 않으면 글자가 깨져 보인다.
# 문서 전체 스타일과 각 run에 eastAsia 글꼴을 함께 박아준다.
_KOREAN_FONT = "맑은 고딕"


def _apply_korean_font(run_or_style) -> None:
    from docx.oxml.ns import qn

    font = run_or_style.font
    font.name = _KOREAN_FONT
    element = run_or_style._element
    rpr = element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), _KOREAN_FONT)
    rfonts.set(qn("w:ascii"), _KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), _KOREAN_FONT)


def _lines_to_docx(lines: list[tuple[str, str]]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    document = Document()

    # 기본 스타일과 제목 스타일 모두에 한글 글꼴을 지정한다.
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "List Bullet"):
        try:
            _apply_korean_font(document.styles[style_name])
        except KeyError:
            continue

    for style, text in lines:
        if style == "spacer":
            document.add_paragraph("")
            continue

        if style in ("title", "heading", "subheading"):
            level = {"title": 0, "heading": 1, "subheading": 2}[style]
            heading = document.add_heading("", level=level)
            run = heading.add_run(text)
        elif style == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(text)
        else:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)

            if style == "meta":
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            elif style == "speaker":
                run.bold = True
            elif style == "translation":
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

        # 모든 run에 한글 글꼴을 직접 지정해야 Word에서 글자가 깨지지 않는다.
        _apply_korean_font(run)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render(lines: list[tuple[str, str]], base_name: str) -> tuple[bytes, str, str]:
    """(내용, 파일명, media type). docx 실패 시 txt로 폴백한다."""
    try:
        return _lines_to_docx(lines), f"{base_name}.docx", DOCX_MEDIA_TYPE
    except Exception:
        logger.exception("docx 생성 실패 — 텍스트 파일로 대체합니다 (%s)", base_name)
        return _lines_to_text(lines), f"{base_name}.txt", TEXT_MEDIA_TYPE


def build_transcript_document(
    session: Session,
    turns: list[Turn],
    *,
    project_title: str | None = None,
) -> tuple[bytes, str, str]:
    lines = _transcript_lines(session, turns, project_title=project_title)
    base_name = f"interview_{safe_filename_part(session.title, fallback=session.id)}"
    return _render(lines, base_name)


# =========================================================
# 프로젝트 리포트
# =========================================================

def _render_content(value: Any, lines: list[tuple[str, str]], *, depth: int = 0) -> None:
    """분석 결과 dict는 스키마가 상황에 따라 달라서 일반적으로 펼친다."""

    if value is None or value == "" or value == [] or value == {}:
        return

    if isinstance(value, dict):
        for key, child in value.items():
            label = str(key).replace("_", " ")
            if isinstance(child, (dict, list)):
                lines.append(("heading" if depth == 0 else "subheading", label))
                _render_content(child, lines, depth=depth + 1)
            else:
                lines.append(("body", f"{label}: {child}"))
        return

    if isinstance(value, list):
        for item in value:
            if isinstance(item, (dict, list)):
                _render_content(item, lines, depth=depth + 1)
                lines.append(("spacer", ""))
            else:
                lines.append(("bullet", str(item)))
        return

    lines.append(("body", str(value)))


def build_project_report_document(
    study: ResearchStudy,
    report: ProjectAggregateReport,
) -> tuple[bytes, str, str]:
    # 1. StudyReportAnalysis 전체 구조가 존재하면 서식/시각화가 반영된 Word 문서 생성 시도
    if report.content and isinstance(report.content, dict):
        try:
            from app.schemas.study_report import StudyReportAnalysis
            from app.export_study_report_word import build_document

            analysis = StudyReportAnalysis.model_validate(report.content)
            document = build_document(analysis)
            buffer = io.BytesIO()
            document.save(buffer)
            base_name = f"project_report_{safe_filename_part(study.title, fallback=study.id)}"
            return buffer.getvalue(), f"{base_name}.docx", DOCX_MEDIA_TYPE
        except Exception:
            logger.exception("Word 시각화 리포트 생성 실패, 기본 텍스트 템플릿으로 폴백")

    # 2. 로컬 모드 또는 부분 데이터일 경우 기본 라인 기반 문서 생성
    lines: list[tuple[str, str]] = [
        ("title", f"프로젝트 리포트 — {study.title}"),
        ("meta", f"연구 목적: {study.research_purpose}"),
        ("meta", f"응답자 수: {report.respondent_count}명"),
        ("meta", f"생성 시각: {_format_datetime(report.generated_at)}"),
        ("meta", f"포함 세션: {len(report.included_session_ids)}건"),
        ("spacer", ""),
    ]

    _render_content(report.content, lines)

    base_name = f"project_report_{safe_filename_part(study.title, fallback=study.id)}"
    return _render(lines, base_name)


XLSX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
JSON_MEDIA_TYPE = "application/json; charset=utf-8"


def _build_fallback_powerbi_tables(
    study: ResearchStudy,
    report: ProjectAggregateReport,
) -> dict[str, list[dict[str, Any]]]:
    from app.export_study_report_bi import EXCEL_TABLE_NAMES

    content = report.content if isinstance(report.content, dict) else {}
    overview_data = content.get("overview") if isinstance(content.get("overview"), dict) else {}
    sessions_data = content.get("sessions") if isinstance(content.get("sessions"), list) else []
    evidence_data = content.get("evidence") if isinstance(content.get("evidence"), list) else []

    # 1. studies
    studies = [
        {
            "study_id": study.id,
            "research_title": overview_data.get("research_title", study.title),
            "research_purpose": overview_data.get("research_purpose", study.research_purpose),
            "core_insight": overview_data.get("core_insight", ""),
            "executive_summary": overview_data.get("executive_summary", ""),
            "participant_count": overview_data.get("participant_count", report.respondent_count),
            "completed_session_count": overview_data.get("completed_session_count", report.respondent_count),
            "question_count": overview_data.get("question_count", len(study.questions)),
        }
    ]

    # 2. participants
    participants: list[dict[str, Any]] = []
    for idx, sess in enumerate(sessions_data, start=1):
        if isinstance(sess, dict):
            p_id = sess.get("respondent_id") or f"P{idx:02d}"
            participants.append(
                {
                    "study_id": study.id,
                    "participant_key": f"{study.id}_{p_id}",
                    "participant_id": p_id,
                    "session_id": sess.get("session_id", ""),
                    "quote_count": sess.get("turn_count", 0),
                }
            )
    if not participants:
        for idx, s_id in enumerate(report.included_session_ids or [], start=1):
            p_id = f"P{idx:02d}"
            participants.append(
                {
                    "study_id": study.id,
                    "participant_key": f"{study.id}_{p_id}",
                    "participant_id": p_id,
                    "session_id": s_id,
                    "quote_count": 0,
                }
            )

    # 3. coverage
    coverage: list[dict[str, Any]] = []
    for q_idx, q in enumerate(study.questions or [], start=1):
        q_text = q.text if hasattr(q, "text") else str(q)
        q_id = q.id if hasattr(q, "id") else f"Q{q_idx:02d}"
        coverage.append(
            {
                "study_id": study.id,
                "coverage_id": f"COV{q_idx:02d}",
                "question_id": q_id,
                "question": q_text,
                "coverage": "not_covered",
                "participant_count": report.respondent_count,
                "covered_participant_count": 0,
                "coverage_rate": 0.0,
                "reason": "",
                "missing_information": "",
            }
        )

    # 4. evidence
    evidence: list[dict[str, Any]] = []
    for e_idx, ev in enumerate(evidence_data, start=1):
        if isinstance(ev, dict):
            p_id = ev.get("respondent_id") or "P01"
            evidence.append(
                {
                    "study_id": study.id,
                    "evidence_key": f"{study.id}_EV{e_idx:03d}",
                    "evidence_id": f"EV{e_idx:03d}",
                    "participant_key": f"{study.id}_{p_id}",
                    "participant_id": p_id,
                    "session_id": ev.get("session_id", ""),
                    "question_id": ev.get("question_id", ""),
                    "quote": ev.get("quote", ""),
                }
            )

    tables: dict[str, list[dict[str, Any]]] = {k: [] for k in EXCEL_TABLE_NAMES.keys()}
    tables["studies"] = studies
    tables["participants"] = participants
    tables["coverage"] = coverage
    tables["evidence"] = evidence
    return tables


def build_powerbi_excel_document(
    study: ResearchStudy,
    report: ProjectAggregateReport,
) -> tuple[bytes, str, str]:
    """Power BI 분석을 위한 정규화된 다중 시트 Excel(.xlsx) 파일을 생성한다."""
    from app.export_study_report_bi import export_powerbi_excel

    base_name = f"study_report_powerbi_{safe_filename_part(study.title, fallback=study.id)}"
    buffer = io.BytesIO()

    tables = None
    if report.content and isinstance(report.content, dict):
        try:
            from app.schemas.study_report import StudyReportAnalysis
            from app.services.report.bi_transformer import get_study_report_bi_transformer

            clean_content = {
                k: v for k, v in report.content.items()
                if k in StudyReportAnalysis.model_fields
            }
            analysis = StudyReportAnalysis.model_validate(clean_content)
            transformer = get_study_report_bi_transformer()
            tables = transformer.transform(analysis)
        except Exception:
            logger.exception("Power BI 테이블 변환 실패, 기본 세션/근거 테이블로 폴백")

    if not tables:
        tables = _build_fallback_powerbi_tables(study, report)

    export_powerbi_excel(tables, buffer)
    return buffer.getvalue(), f"{base_name}.xlsx", XLSX_MEDIA_TYPE


def build_project_report_json(
    study: ResearchStudy,
    report: ProjectAggregateReport,
) -> tuple[bytes, str, str]:
    """분석 원본 전체 데이터를 JSON 파일로 내려준다."""
    import json

    base_name = f"project_report_{safe_filename_part(study.title, fallback=study.id)}"
    content_bytes = json.dumps(
        report.content or {},
        ensure_ascii=False,
        indent=2,
    ).encode("utf-8")
    return content_bytes, f"{base_name}.json", JSON_MEDIA_TYPE
