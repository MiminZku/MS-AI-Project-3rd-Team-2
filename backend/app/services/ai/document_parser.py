"""워드(.docx), PDF, 마크다운(.md) 인터뷰 가이드라인 문서로부터
구조화된 질문 트리(QuestionNode)와 리서치 개요를 자동 추출하는 AI 문서 파서.
"""

from __future__ import annotations

import io
import json
import logging
import re
from typing import Literal
from pydantic import BaseModel, Field

from app.core.config import get_settings
from app.schemas.session import QuestionNode

logger = logging.getLogger(__name__)


class ExtractedQuestion(BaseModel):
    order: int
    text: str
    intent: str = Field(description="이 질문을 하는 리서치 의도/목적")
    keywords: list[str] = Field(default_factory=list, description="관련 핵심 키워드 목록")
    branches: dict[str, str] = Field(default_factory=dict, description="응답 갈래별(예: 긍정, 부정 등) 파생질문")


class ParsedGuideResult(BaseModel):
    title: str = Field(description="추출된 리서치/인터뷰 제목")
    research_purpose: str = Field(description="추출된 리서치 목적 및 배경")
    target_screening: str = Field(default="", description="대상자 선정/스크리닝 기준")
    questions: list[ExtractedQuestion]


class DocumentParser:
    def __init__(self) -> None:
        self.settings = get_settings()

    def extract_text_from_bytes(self, content: bytes, filename: str) -> str:
        """파일 확장자에 맞게 텍스트를 추출."""
        ext = filename.lower().split(".")[-1] if "." in filename else ""

        if ext in ("docx", "doc"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            full_text.append(" | ".join(row_text))
                return "\n\n".join(full_text)
            except Exception as e:
                logger.exception("Word docx 텍스트 추출 실패: %s", e)
                return content.decode("utf-8", errors="ignore")

        elif ext == "pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content))
                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                return "\n\n".join(pages)
            except Exception as e:
                logger.exception("PDF 텍스트 추출 실패: %s", e)
                return content.decode("utf-8", errors="ignore")

        else:
            # md, txt 등 일반 텍스트
            return content.decode("utf-8", errors="ignore")

    async def parse_guide(self, raw_text: str) -> ParsedGuideResult:
        """LLM을 이용해 비구조화된 인터뷰 지침서/가이드라인 문서를 정형화된 JSON 질문 트리로 파싱."""
        if not self.settings.use_azure_openai:
            # Stub/Rule-based Fallback
            return self._fallback_rule_parse(raw_text)

        from openai import AsyncAzureOpenAI
        client = AsyncAzureOpenAI(
            azure_endpoint=self.settings.azure_openai_endpoint,
            api_key=self.settings.azure_openai_api_key,
            api_version=self.settings.azure_openai_api_version,
            timeout=90.0,
        )

        system_prompt = """당신은 세계 최고 수준의 사용자 경험(UX) 리서치 및 인터뷰 설계 전문가입니다.
사용자가 업로드한 인터뷰 가이드라인(워드, PDF, 마크다운 등) 원문 텍스트를 정밀 분석하여,
AI 인터뷰어가 실시간 인터뷰에서 활용할 구조화된 질문 트리 JSON을 생성하세요.

[추출 및 정제 지침]
1. title: 인터뷰 주제/제목 (원문에서 파악된 주제를 명확하고 간결하게 작성)
2. research_purpose: 이 인터뷰를 통해 밝혀내고자 하는 핵심 조사 목적 및 문제의식
3. target_screening: 참가자 선정 조건 (없으면 빈 문자열)
4. questions: 인터뷰에서 진행할 핵심 질문 목록 (순서대로)
   - order: 1부터 시작하는 순번
   - text: 인터뷰이가 편안하게 답변할 수 있도록 정제된 실제 질문 문장 (구어체 친화적)
   - intent: 이 질문을 던지는 리서처의 핵심 의도 (무엇을 알아내려 하는가)
   - keywords: 질문과 관련된 핵심 키워드 3~5개
   - branches: 인터뷰이의 답변 내용이나 상황 조건에 따라 파생될 꼬리질문 딕셔너리 (키-값 쌍)
     ⚠️ 중요: branches의 key는 "탐색 갈래 1" 같은 모호한 이름이 아니라, **"어떤 답변이나 상황일 때 이 질문을 하는지" 구체적인 조건/상황/관점**을 명시하세요.
"""

        try:
            response = await client.chat.completions.create(
                model=self.settings.azure_openai_chat_deployment,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"[가이드라인 원문]\n\n{raw_text}"},
                ],
                response_format={"type": "json_object"},
                temperature=0.2,
            )
            content = response.choices[0].message.content or "{}"
            data = json.loads(content)
            return ParsedGuideResult.model_validate(data)
        except Exception as e:
            logger.exception("LLM 질문 가이드 파싱 실패 (룰 기반 파서로 대체): %s", e)
            return self._fallback_rule_parse(raw_text)

    def _fallback_rule_parse(self, text: str) -> ParsedGuideResult:
        """Azure OpenAI 호출 실패 또는 오프라인 시 문서 원문에서 동적으로 제목, 목적, 질문을 추출하는 룰 기반 파서."""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        
        # 1. 제목 동적 추출
        title = "사용자 리서치 인터뷰"
        title_match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
        if not title_match:
            title_match = re.search(r"(?:인터뷰\s*주제|프로젝트명|제목)\s*[:：]\s*(.+)", text)
        if not title_match:
            title_match = re.search(r"^##\s+(.+)$", text, re.MULTILINE)
        
        if title_match:
            title = title_match.group(1).strip(" #*\"'")
        elif lines:
            # 첫 번째 의미 있는 줄을 제목으로 사용
            candidate = lines[0].strip(" #*\"'-")
            if len(candidate) >= 3:
                title = candidate

        # 2. 조사 목적 동적 추출
        purpose = f"{title} 관련 사용자 경험 및 심층 의견 파악"
        purpose_match = re.search(r"(?:조사\s*목적|목적|연구\s*목적|배경)\s*[:：]\s*(.+?)(?=\n\*|\n---|\n##|\n\n|$)", text, re.DOTALL)
        if purpose_match:
            purpose = purpose_match.group(1).strip(" *\"'\n\r")

        # 3. 질문 및 파생질문 동적 추출
        questions: list[ExtractedQuestion] = []
        
        # 방식 A: Markdown Section 포맷 (### [Section N])
        sections = re.split(r"###\s*\[Section\s*\d+\]", text)
        if len(sections) > 1:
            order = 1
            for sec in sections[1:]:
                goal_match = re.search(r"\*\s*\*\*목표:\*\*\s*(.+)", sec)
                goal_text = goal_match.group(1).strip() if goal_match else "핵심 현황 파악"
                
                core_q_match = re.search(r"\*\s*\*\*핵심 질문:\*\*\s*\n\s*\*\s*[\"\'\*]*(.+?)[\"\'\*]*\s*$", sec, re.MULTILINE)
                if not core_q_match:
                    core_q_match = re.search(r"\*\s*\*\*핵심 질문:\*\*\s*\n\s*\*\s*(.+)", sec)
                
                if core_q_match:
                    q_text = core_q_match.group(1).strip(" *\"'")
                    branches: dict[str, str] = {}
                    probing_part = sec.split("Probing (탐색 갈래):")[-1] if "Probing (탐색 갈래):" in sec else ""
                    probing_lines = [l.strip() for l in probing_part.splitlines() if l.strip().startswith("*")]
                    
                    for pl in probing_lines:
                        pl_clean = re.sub(r"^\*\s*", "", pl).strip()
                        if not pl_clean or pl_clean == "*":
                            continue
                        cond_match = re.match(r"^\*?\*?\((.+?)\):?\*?\*?\s*(.+)", pl_clean)
                        if cond_match:
                            branches[cond_match.group(1).strip()] = cond_match.group(2).strip()
                        else:
                            clean_branch_q = pl_clean.strip(" *\"'")
                            branches[f"상세 경험 ({clean_branch_q[:12]}...)"] = clean_branch_q

                    questions.append(ExtractedQuestion(
                        order=order,
                        text=q_text,
                        intent=goal_text,
                        keywords=[title[:10], "경험", "사용성"],
                        branches=branches,
                    ))
                    order += 1

        # 방식 B: 번호 매겨진 질문 목록 (1. ..., 2. ...) 또는 Q1, Q2
        if not questions:
            q_matches = re.findall(r"(?:^|\n)(?:Q?\s*(\d+)[\.\)]|\*\s*(\d+)[\.\)])\s*(.+)", text)
            if q_matches:
                for idx, match in enumerate(q_matches, start=1):
                    q_num = match[0] or match[1] or str(idx)
                    q_content = match[2].strip(" *\"'")
                    if len(q_content) >= 5:
                        questions.append(ExtractedQuestion(
                            order=int(q_num) if q_num.isdigit() else idx,
                            text=q_content,
                            intent=f"{title} 관련 {idx}번 핵심 확인",
                            keywords=[title[:10], f"항목{idx}"],
                            branches={},
                        ))

        # 방식 C: 최종 fallback
        if not questions:
            for idx, line in enumerate(lines[:5], start=1):
                clean_line = line.strip(" #*\"'-")
                if len(clean_line) >= 8 and any(clean_line.endswith(end) for end in ("?", "요", "까", "음")):
                    questions.append(ExtractedQuestion(
                        order=idx,
                        text=clean_line,
                        intent=f"{title} 관련 질문",
                        keywords=[title[:10]],
                        branches={},
                    ))

        if not questions:
            questions = [
                ExtractedQuestion(
                    order=1,
                    text=f"안녕하세요, {title}에 대한 경험이나 생각에 대해 편하게 말씀해 주시겠어요?",
                    intent="초기 현황 파악 및 라포 형성",
                    keywords=[title[:10], "기본경험"],
                    branches={},
                )
            ]

        return ParsedGuideResult(
            title=title,
            research_purpose=purpose,
            target_screening="해당 주제 관련 실사용자 또는 관심 대상자",
            questions=questions,
        )


_parser: DocumentParser | None = None

def get_document_parser() -> DocumentParser:
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
