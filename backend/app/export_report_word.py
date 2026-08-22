from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# =========================================================
# 경로
# =========================================================

BASE_DIR = Path(__file__).resolve().parent

REPORT_DIR = BASE_DIR / "ai-interview-report"

INPUT_JSON = (
    REPORT_DIR
    / "participant_report.json"
)

OUTPUT_DOCX = (
    REPORT_DIR
    / "participant_report.docx"
)


# =========================================================
# Word 기본 설정
# =========================================================

FONT_NAME = "맑은 고딕"

COVERAGE_LABEL = {
    "high": "충분",
    "medium": "보통",
    "low": "부족",
    "not_covered": "미확인",
}

STRENGTH_LABEL = {
    "strong": "강함",
    "moderate": "보통",
    "weak": "약함",
}

PRIORITY_LABEL = {
    "high": "높음",
    "medium": "보통",
    "low": "낮음",
}

SEVERITY_LABEL = {
    "high": "높음",
    "medium": "보통",
    "low": "낮음",
}

SOURCE_LABEL = {
    "explicit_user_request": "응답자 직접 요구",
    "derived_opportunity": "AI 분석 도출",
}

OPPORTUNITY_TYPE_LABEL = {
    "product": "제품",
    "service": "서비스",
    "process": "프로세스",
    "policy": "정책",
    "communication": "커뮤니케이션",
    "research": "추가 연구",
    "other": "기타",
}


# =========================================================
# 공통 Word Helper
# =========================================================

def set_run_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
) -> None:

    run.font.name = FONT_NAME

    if size is not None:
        run.font.size = Pt(size)

    if bold is not None:
        run.bold = bold

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        FONT_NAME,
    )


def set_paragraph_font(
    paragraph,
    size: float = 10,
) -> None:

    for run in paragraph.runs:
        set_run_font(
            run,
            size=size,
        )


def set_cell_background(
    cell,
    fill: str,
) -> None:

    tc_pr = cell._tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(
        qn("w:fill"),
        fill,
    )

    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    size: float = 9,
) -> None:

    cell.text = ""

    paragraph = cell.paragraphs[0]

    run = paragraph.add_run(
        str(text)
    )

    set_run_font(
        run,
        size=size,
        bold=bold,
    )

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


def add_title(
    document: Document,
    text: str,
) -> None:

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.space_after = Pt(10)

    run = paragraph.add_run(text)

    set_run_font(
        run,
        size=26,
        bold=True,
    )


def add_subtitle(
    document: Document,
    text: str,
) -> None:

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(text)

    set_run_font(
        run,
        size=12,
    )


def add_heading(
    document: Document,
    text: str,
    level: int = 1,
) -> None:

    paragraph = document.add_heading(
        level=level
    )

    run = paragraph.add_run(text)

    set_run_font(
        run,
        size=(
            17
            if level == 1
            else 13
        ),
        bold=True,
    )


def add_body(
    document: Document,
    text: str,
    bold_prefix: str | None = None,
) -> None:

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_after = (
        Pt(6)
    )

    paragraph.paragraph_format.line_spacing = (
        1.3
    )

    if bold_prefix:

        prefix_run = paragraph.add_run(
            bold_prefix
        )

        set_run_font(
            prefix_run,
            size=10,
            bold=True,
        )

    run = paragraph.add_run(
        str(text)
    )

    set_run_font(
        run,
        size=10,
    )


def add_bullet(
    document: Document,
    text: str,
) -> None:

    paragraph = document.add_paragraph(
        style="List Bullet"
    )

    run = paragraph.add_run(
        str(text)
    )

    set_run_font(
        run,
        size=10,
    )


def add_evidence_ids(
    document: Document,
    evidence_ids: list[str],
) -> None:

    if not evidence_ids:
        return

    add_body(
        document,
        ", ".join(evidence_ids),
        bold_prefix="근거: ",
    )


def add_separator(
    document: Document,
) -> None:

    paragraph = document.add_paragraph()

    p_pr = paragraph._p.get_or_add_pPr()

    border = OxmlElement(
        "w:pBdr"
    )

    bottom = OxmlElement(
        "w:bottom"
    )

    bottom.set(
        qn("w:val"),
        "single",
    )

    bottom.set(
        qn("w:sz"),
        "4",
    )

    bottom.set(
        qn("w:space"),
        "1",
    )

    bottom.set(
        qn("w:color"),
        "D9D9D9",
    )

    border.append(bottom)
    p_pr.append(border)


# =========================================================
# 1. 표지
# =========================================================

def add_cover(
    document: Document,
    report: dict[str, Any],
) -> None:

    data = report["data"]
    metadata = data.get(
        "metadata",
        {},
    )

    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()

    add_title(
        document,
        "Individual Interview Report",
    )

    title = metadata.get(
        "research_title",
        "인터뷰 리포트",
    )

    add_subtitle(
        document,
        title,
    )

    document.add_paragraph()
    document.add_paragraph()

    info_table = document.add_table(
        rows=4,
        cols=2,
    )

    info_table.style = (
        "Table Grid"
    )

    rows = [
        (
            "Session ID",
            report.get(
                "session_id",
                "-",
            ),
        ),
        (
            "Study ID",
            metadata.get(
                "study_id",
                "-",
            ),
        ),
        (
            "질문 수",
            metadata.get(
                "question_count",
                "-",
            ),
        ),
        (
            "Information Slot 수",
            metadata.get(
                "information_slot_count",
                "-",
            ),
        ),
    ]

    for row, values in zip(
        info_table.rows,
        rows,
    ):

        set_cell_background(
            row.cells[0],
            "F2F2F2",
        )

        set_cell_text(
            row.cells[0],
            values[0],
            bold=True,
        )

        set_cell_text(
            row.cells[1],
            values[1],
        )

    document.add_page_break()


# =========================================================
# 2. Executive Summary
# =========================================================

def add_executive_summary(
    document: Document,
    data: dict[str, Any],
) -> None:

    summary = data.get(
        "executive_summary",
        {},
    )

    add_heading(
        document,
        "1. Executive Summary",
    )

    add_body(
        document,
        summary.get(
            "core_insight",
            "",
        ),
        bold_prefix="핵심 결론: ",
    )

    add_body(
        document,
        summary.get(
            "summary",
            "",
        ),
        bold_prefix="요약: ",
    )

    takeaways = summary.get(
        "key_takeaways",
        [],
    )

    if takeaways:

        add_heading(
            document,
            "핵심 Takeaways",
            level=2,
        )

        for item in takeaways:

            add_bullet(
                document,
                item.get(
                    "point",
                    "",
                ),
            )

            add_evidence_ids(
                document,
                item.get(
                    "evidence_ids",
                    [],
                ),
            )


# =========================================================
# 3. Participant Context
# =========================================================

def add_participant_context(
    document: Document,
    data: dict[str, Any],
) -> None:

    context = data.get(
        "participant_context",
        {},
    )

    add_heading(
        document,
        "2. Participant Context",
    )

    add_body(
        document,
        context.get(
            "summary",
            "",
        ),
    )

    attributes = context.get(
        "attributes",
        [],
    )

    if not attributes:
        return

    table = document.add_table(
        rows=1,
        cols=3,
    )

    table.style = "Table Grid"

    headers = [
        "항목",
        "내용",
        "Evidence",
    ]

    for index, header in enumerate(
        headers
    ):

        set_cell_background(
            table.rows[0].cells[index],
            "D9EAF7",
        )

        set_cell_text(
            table.rows[0].cells[index],
            header,
            bold=True,
        )

    for item in attributes:

        cells = table.add_row().cells

        set_cell_text(
            cells[0],
            item.get(
                "name",
                "",
            ),
            bold=True,
        )

        set_cell_text(
            cells[1],
            item.get(
                "value",
                "",
            ),
        )

        set_cell_text(
            cells[2],
            ", ".join(
                item.get(
                    "evidence_ids",
                    [],
                )
            ),
        )


# =========================================================
# 4. Research Coverage
# =========================================================

def add_research_coverage(
    document: Document,
    data: dict[str, Any],
) -> None:

    coverage = data.get(
        "research_coverage",
        {},
    )

    add_heading(
        document,
        "3. Question Coverage",
    )

    overall = coverage.get(
        "overall_coverage",
        "-",
    )

    add_body(
        document,
        COVERAGE_LABEL.get(
            overall,
            overall,
        ),
        bold_prefix="전체 Coverage: ",
    )

    for item in coverage.get(
        "items",
        [],
    ):

        question_id = item.get(
            "question_id",
            "",
        )

        question = item.get(
            "question",
            "",
        )

        add_heading(
            document,
            f"{question_id.upper()} · {question}",
            level=2,
        )

        coverage_value = item.get(
            "coverage",
            "",
        )

        add_body(
            document,
            COVERAGE_LABEL.get(
                coverage_value,
                coverage_value,
            ),
            bold_prefix="확보 수준: ",
        )

        add_body(
            document,
            item.get(
                "reason",
                "",
            ),
            bold_prefix="판단 이유: ",
        )

        missing = item.get(
            "missing_information",
            [],
        )

        if missing:

            add_body(
                document,
                "",
                bold_prefix="추가 확인 필요:",
            )

            for missing_item in missing:
                add_bullet(
                    document,
                    missing_item,
                )

        add_evidence_ids(
            document,
            item.get(
                "evidence_ids",
                [],
            ),
        )

        add_separator(document)


# =========================================================
# 5. Slot Coverage
# =========================================================

def add_slot_coverage(
    document: Document,
    data: dict[str, Any],
) -> None:

    slot_coverage = data.get(
        "slot_coverage",
        {},
    )

    add_heading(
        document,
        "4. Information Slot Coverage",
    )

    overall = slot_coverage.get(
        "overall_coverage",
        "-",
    )

    add_body(
        document,
        COVERAGE_LABEL.get(
            overall,
            overall,
        ),
        bold_prefix="전체 Slot Coverage: ",
    )

    table = document.add_table(
        rows=1,
        cols=5,
    )

    table.style = "Table Grid"

    headers = [
        "Question",
        "Slot",
        "확보 수준",
        "판단 이유",
        "Evidence",
    ]

    for i, header in enumerate(
        headers
    ):

        set_cell_background(
            table.rows[0].cells[i],
            "D9EAF7",
        )

        set_cell_text(
            table.rows[0].cells[i],
            header,
            bold=True,
            size=8,
        )

    for item in slot_coverage.get(
        "items",
        [],
    ):

        cells = table.add_row().cells

        coverage_value = item.get(
            "coverage",
            "",
        )

        values = [
            item.get(
                "question_id",
                "",
            ),
            item.get(
                "slot_name",
                "",
            ),
            COVERAGE_LABEL.get(
                coverage_value,
                coverage_value,
            ),
            item.get(
                "reason",
                "",
            ),
            ", ".join(
                item.get(
                    "evidence_ids",
                    [],
                )
            ),
        ]

        for i, value in enumerate(
            values
        ):

            set_cell_text(
                cells[i],
                value,
                size=8,
            )

    document.add_paragraph()

    add_heading(
        document,
        "미확인 / 추가 확인 정보",
        level=2,
    )

    for item in slot_coverage.get(
        "items",
        [],
    ):

        missing = item.get(
            "missing_information",
            [],
        )

        if not missing:
            continue

        add_body(
            document,
            item.get(
                "slot_name",
                "",
            ),
            bold_prefix="Slot: ",
        )

        for missing_item in missing:

            add_bullet(
                document,
                missing_item,
            )


# =========================================================
# 6. Key Findings
# =========================================================

def add_key_findings(
    document: Document,
    data: dict[str, Any],
) -> None:

    add_heading(
        document,
        "5. Key Findings",
    )

    findings = data.get(
        "key_findings",
        [],
    )

    if not findings:

        add_body(
            document,
            "핵심 Finding이 확인되지 않았습니다.",
        )

        return

    for index, item in enumerate(
        findings,
        start=1,
    ):

        add_heading(
            document,
            (
                f"{index}. "
                f"{item.get('title', '')}"
            ),
            level=2,
        )

        add_body(
            document,
            item.get(
                "summary",
                "",
            ),
        )

        strength = item.get(
            "evidence_strength",
            "",
        )

        add_body(
            document,
            STRENGTH_LABEL.get(
                strength,
                strength,
            ),
            bold_prefix="Evidence 강도: ",
        )

        add_evidence_ids(
            document,
            item.get(
                "evidence_ids",
                [],
            ),
        )


# =========================================================
# 7. Themes
# =========================================================

def add_themes(
    document: Document,
    data: dict[str, Any],
) -> None:

    add_heading(
        document,
        "6. Themes",
    )

    themes = data.get(
        "themes",
        [],
    )

    if not themes:

        add_body(
            document,
            "별도의 상위 Theme이 확인되지 않았습니다.",
        )

        return

    for item in themes:

        add_heading(
            document,
            item.get(
                "theme",
                "",
            ),
            level=2,
        )

        add_body(
            document,
            item.get(
                "description",
                "",
            ),
        )

        add_evidence_ids(
            document,
            item.get(
                "evidence_ids",
                [],
            ),
        )


# =========================================================
# 8. Key Drivers
# =========================================================

def add_key_drivers(
    document: Document,
    data: dict[str, Any],
) -> None:

    add_heading(
        document,
        "7. Key Drivers",
    )

    drivers = data.get(
        "key_drivers",
        [],
    )

    if not drivers:

        add_body(
            document,
            "명확한 Driver가 확인되지 않았습니다.",
        )

        return

    for item in drivers:

        add_heading(
            document,
            item.get(
                "driver",
                "",
            ),
            level=2,
        )

        strength = item.get(
            "strength",
            "",
        )

        add_body(
            document,
            PRIORITY_LABEL.get(
                strength,
                strength,
            ),
            bold_prefix="영향도: ",
        )

        add_body(
            document,
            item.get(
                "description",
                "",
            ),
        )

        add_evidence_ids(
            document,
            item.get(
                "evidence_ids",
                [],
            ),
        )


# =========================================================
# 9. Needs & Pain Points
# =========================================================

def add_needs_and_pain_points(
    document: Document,
    data: dict[str, Any],
) -> None:

    add_heading(
        document,
        "8. Needs & Pain Points",
    )

    items = data.get(
        "needs_and_pain_points",
        [],
    )

    if not items:

        add_body(
            document,
            "확인된 Need 또는 Pain Point가 없습니다.",
        )

        return

    for item in items:

        type_value = item.get(
            "type",
            "",
        )

        type_label = (
            "Pain Point"
            if type_value == "pain_point"
            else "Need"
        )

        add_heading(
            document,
            (
                f"[{type_label}] "
                f"{item.get('title', '')}"
            ),
            level=2,
        )

        severity = item.get(
            "severity",
            "",
        )

        add_body(
            document,
            SEVERITY_LABEL.get(
                severity,
                severity,
            ),
            bold_prefix="중요도: ",
        )

        add_body(
            document,
            item.get(
                "situation",
                "",
            ),
            bold_prefix="상황: ",
        )

        add_body(
            document,
            item.get(
                "impact",
                "",
            ),
            bold_prefix="영향: ",
        )

        add_evidence_ids(
            document,
            item.get(
                "evidence_ids",
                [],
            ),
        )


# =========================================================
# 10. Decision Dynamics
# =========================================================

def add_decision_dynamics(
    document: Document,
    data: dict[str, Any],
) -> None:

    add_heading(
        document,
        "9. Decision / Behavior Dynamics",
    )

    decision = data.get(
        "decision_dynamics"
    )

    if decision is None:

        add_body(
            document,
            (
                "이번 인터뷰에서는 미래의 구매, 전환, "
                "이탈 또는 재사용 의도를 판단할 수 있는 "
                "충분한 Evidence가 확보되지 않았습니다."
            ),
        )

        return

    add_body(
        document,
        decision.get(
            "current_state",
            "",
        ),
        bold_prefix="현재 상태: ",
    )

    signal = decision.get(
        "behavioral_signal",
        "",
    )

    add_body(
        document,
        signal,
        bold_prefix="Behavioral Signal: ",
    )

    sections = [
        (
            "의사결정 요인",
            "decision_factors",
            "factor",
        ),
        (
            "Barrier",
            "barriers",
            "barrier",
        ),
        (
            "Trigger",
            "triggers",
            "trigger",
        ),
    ]

    for title, key, text_key in sections:

        items = decision.get(
            key,
            [],
        )

        if not items:
            continue

        add_heading(
            document,
            title,
            level=2,
        )

        for item in items:

            add_bullet(
                document,
                item.get(
                    text_key,
                    "",
                ),
            )

            add_evidence_ids(
                document,
                item.get(
                    "evidence_ids",
                    [],
                ),
            )


# =========================================================
# 11. Opportunities
# =========================================================

def add_opportunities(
    document: Document,
    data: dict[str, Any],
) -> None:

    add_heading(
        document,
        "10. Opportunities",
    )

    opportunities = data.get(
        "opportunities",
        [],
    )

    if not opportunities:

        add_body(
            document,
            "현재 Evidence에서 도출된 Opportunity가 없습니다.",
        )

        return

    for index, item in enumerate(
        opportunities,
        start=1,
    ):

        add_heading(
            document,
            (
                f"{index}. "
                f"{item.get('opportunity', '')}"
            ),
            level=2,
        )

        source_type = item.get(
            "source_type",
            "",
        )

        opportunity_type = item.get(
            "opportunity_type",
            "",
        )

        priority = item.get(
            "priority",
            "",
        )

        add_body(
            document,
            SOURCE_LABEL.get(
                source_type,
                source_type,
            ),
            bold_prefix="출처: ",
        )

        add_body(
            document,
            OPPORTUNITY_TYPE_LABEL.get(
                opportunity_type,
                opportunity_type,
            ),
            bold_prefix="유형: ",
        )

        add_body(
            document,
            PRIORITY_LABEL.get(
                priority,
                priority,
            ),
            bold_prefix="우선순위: ",
        )

        add_body(
            document,
            item.get(
                "problem_or_need",
                "",
            ),
            bold_prefix="문제 / Need: ",
        )

        add_body(
            document,
            item.get(
                "expected_value",
                "",
            ),
            bold_prefix="기대 가치: ",
        )

        add_evidence_ids(
            document,
            item.get(
                "evidence_ids",
                [],
            ),
        )


# =========================================================
# 12. Observer Intervention
# =========================================================

def add_observer_analysis(
    document: Document,
    data: dict[str, Any],
) -> None:

    add_heading(
        document,
        "11. Observer Intervention",
    )

    items = data.get(
        "observer_intervention_analysis",
        [],
    )

    if not items:

        add_body(
            document,
            "Observer 개입이 없습니다.",
        )

        return

    for item in items:

        add_body(
            document,
            item.get(
                "instruction",
                "",
            ),
            bold_prefix="Observer 지시: ",
        )

        add_body(
            document,
            str(
                item.get(
                    "applied_turn",
                    "",
                )
            ),
            bold_prefix="적용 Turn: ",
        )

        value = item.get(
            "research_value",
            "",
        )

        add_body(
            document,
            PRIORITY_LABEL.get(
                value,
                value,
            ),
            bold_prefix="Research Value: ",
        )

        add_body(
            document,
            item.get(
                "impact",
                "",
            ),
            bold_prefix="확보된 가치: ",
        )

        add_evidence_ids(
            document,
            item.get(
                "resulting_evidence_ids",
                [],
            ),
        )


# =========================================================
# 13. Researcher Attention
# =========================================================

def add_researcher_attention(
    document: Document,
    data: dict[str, Any],
) -> None:

    add_heading(
        document,
        "12. Researcher Attention",
    )

    items = data.get(
        "researcher_attention",
        [],
    )

    if not items:

        add_body(
            document,
            "추가 확인이 필요한 핵심 항목이 없습니다.",
        )

        return

    for item in items:

        priority = item.get(
            "priority",
            "",
        )

        add_heading(
            document,
            item.get(
                "topic",
                "",
            ),
            level=2,
        )

        add_body(
            document,
            PRIORITY_LABEL.get(
                priority,
                priority,
            ),
            bold_prefix="우선순위: ",
        )

        add_body(
            document,
            item.get(
                "reason",
                "",
            ),
            bold_prefix="추가 확인 이유: ",
        )


# =========================================================
# 14. Evidence Library
# =========================================================

def add_evidence_library(
    document: Document,
    data: dict[str, Any],
) -> None:

    add_heading(
        document,
        "13. Evidence Library",
    )

    evidence_items = data.get(
        "evidence",
        [],
    )

    if not evidence_items:

        add_body(
            document,
            "Evidence가 없습니다.",
        )

        return

    for item in evidence_items:

        evidence_id = item.get(
            "evidence_id",
            "",
        )

        turn_index = item.get(
            "turn_index",
            "",
        )

        add_heading(
            document,
            (
                f"{evidence_id} "
                f"· Turn {turn_index}"
            ),
            level=2,
        )

        add_body(
            document,
            item.get(
                "quote",
                "",
            ),
        )


# =========================================================
# 15. Metadata
# =========================================================

def add_metadata(
    document: Document,
    data: dict[str, Any],
) -> None:

    metadata = data.get(
        "metadata",
        {},
    )

    add_heading(
        document,
        "14. Metadata",
    )

    rows = [
        (
            "Study ID",
            metadata.get(
                "study_id",
                "-",
            ),
        ),
        (
            "Research Title",
            metadata.get(
                "research_title",
                "-",
            ),
        ),
        (
            "Interview Duration",
            (
                f"{metadata.get('duration_minutes', '-')}분"
            ),
        ),
        (
            "Question Count",
            metadata.get(
                "question_count",
                "-",
            ),
        ),
        (
            "Information Slot Count",
            metadata.get(
                "information_slot_count",
                "-",
            ),
        ),
        (
            "Turn Count",
            metadata.get(
                "turn_count",
                "-",
            ),
        ),
        (
            "Interviewee Turn Count",
            metadata.get(
                "interviewee_turn_count",
                "-",
            ),
        ),
        (
            "Observer Instruction Count",
            metadata.get(
                "instruction_count",
                "-",
            ),
        ),
    ]

    table = document.add_table(
        rows=0,
        cols=2,
    )

    table.style = "Table Grid"

    for name, value in rows:

        cells = table.add_row().cells

        set_cell_background(
            cells[0],
            "F2F2F2",
        )

        set_cell_text(
            cells[0],
            name,
            bold=True,
        )

        set_cell_text(
            cells[1],
            str(value),
        )


# =========================================================
# Word 전체 스타일
# =========================================================

def configure_document(
    document: Document,
) -> None:

    section = document.sections[0]

    section.top_margin = Inches(
        0.7
    )

    section.bottom_margin = Inches(
        0.7
    )

    section.left_margin = Inches(
        0.75
    )

    section.right_margin = Inches(
        0.75
    )

    normal_style = (
        document.styles["Normal"]
    )

    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(10)

    normal_style._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        FONT_NAME,
    )

    for style_name in [
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
    ]:

        style = document.styles[
            style_name
        ]

        style.font.name = FONT_NAME

        style._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            FONT_NAME,
        )


# =========================================================
# Main
# =========================================================

def export_report_to_word(
    input_json: Path,
    output_docx: Path,
) -> Path:

    if not input_json.exists():

        raise FileNotFoundError(
            f"리포트 JSON을 찾을 수 없습니다: "
            f"{input_json}"
        )

    with input_json.open(
        "r",
        encoding="utf-8",
    ) as file:

        report = json.load(file)

    data = report.get(
        "data",
        {},
    )

    document = Document()

    configure_document(
        document
    )

    add_cover(
        document,
        report,
    )

    add_executive_summary(
        document,
        data,
    )

    add_participant_context(
        document,
        data,
    )

    add_research_coverage(
        document,
        data,
    )

    add_slot_coverage(
        document,
        data,
    )

    add_key_findings(
        document,
        data,
    )

    add_themes(
        document,
        data,
    )

    add_key_drivers(
        document,
        data,
    )

    add_needs_and_pain_points(
        document,
        data,
    )

    add_decision_dynamics(
        document,
        data,
    )

    add_opportunities(
        document,
        data,
    )

    add_observer_analysis(
        document,
        data,
    )

    add_researcher_attention(
        document,
        data,
    )

    add_evidence_library(
        document,
        data,
    )

    add_metadata(
        document,
        data,
    )

    output_docx.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        output_docx
    )

    return output_docx


def main() -> None:

    output_path = (
        export_report_to_word(
            input_json=INPUT_JSON,
            output_docx=OUTPUT_DOCX,
        )
    )

    print()
    print(
        "========================================"
    )
    print(
        "Word 리포트 생성 완료"
    )
    print(
        "========================================"
    )
    print(
        f"입력 JSON: {INPUT_JSON}"
    )
    print(
        f"출력 Word: {output_path}"
    )


if __name__ == "__main__":
    main()